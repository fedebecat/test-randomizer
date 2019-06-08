from docx import Document
import random
from tqdm import tqdm

filename = 'test.docx'
num_copies = 10
randomize_questions = True
randomize_answers = True

document = Document(filename)
paragraphs = document.paragraphs
header = document.sections[0].header
header_text = ''.join([x.text for x in header.paragraphs])

''' Read doc '''
questions = {}
cur_question = None
for e, p in enumerate(paragraphs):
    line = ''.join([x.text for x in p.runs])
    if len(line) == 0:
        continue
    if (all([x.bold for x in p.runs])):
        cur_question = line
        questions[cur_question] = []
    else:
        questions[cur_question].append(line)

''' Create new doc '''
doc = Document()
doc.sections[0].header.paragraphs[0].text = header_text

print('Generating copies...')
for copy in tqdm(range(num_copies)):
    all_questions = list(questions.keys())
    if randomize_questions:
        random.shuffle(all_questions)
    question_counter = 0
    for e, k in enumerate(all_questions):
        if len(questions[k]) == 0:
            continue
        p = doc.add_paragraph()
        runner = p.add_run('{}. {}'.format(question_counter+1, k))
        runner.bold = True
        answers = questions[k]
        if randomize_answers:
            random.shuffle(answers)
        for v in answers:
            if sum([x == '_' for x in v])/len(v) > 0.8:
                doc.add_paragraph(
                    '{}'.format('_'*100)
                )
            else:
                doc.add_paragraph(
                    v, style='List Bullet 2'
                )
        question_counter += 1
        if question_counter == 6:
            doc.add_page_break()

    doc.add_page_break()
doc.save('output.docx')